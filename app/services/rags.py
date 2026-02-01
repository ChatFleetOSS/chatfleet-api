# generated-by: codex-agent 2025-02-15T00:20:00Z
"""
RAG persistence, document lifecycle management, and access control helpers.
"""

from __future__ import annotations

import asyncio
import logging
import hashlib
import chardet
import json
import random
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple
from uuid import uuid4

from bson import ObjectId
from fastapi import UploadFile, status
from motor.motor_asyncio import AsyncIOMotorCollection
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from odf.opendocument import load as load_odt
from odf import teletype
from odf.text import P as OdtParagraph
from odf.table import Table as OdtTable, TableRow as OdtTableRow, TableCell as OdtTableCell
from pymongo.errors import DuplicateKeyError

from app.core.config import settings
from app.core.database import get_collection as mongo_collection
from app.models.rag import (
    IndexStatusEnum,
    IndexStatusResponse,
    RagDoc,
    RagDocsResponse,
    RagSummary,
    RagUploadAccepted,
    RagCreateRequest,
    RagUser,
    RagUsersResponse,
)
from app.models.jobs import JobProgressTotals
from app.services.chunker import ChunkWithPage, PageText, chunk_pdf
from app.services.embeddings import embed_texts
from app.services.jobs import JobRecord, job_manager
from app.services.logging import write_system_log
from app.services.users import find_user_by_id, update_user_rags
from app.services.vectorstore import build_index as rebuild_vector_index
from app.services.vectorstore import persist_doc_payload, remove_doc_payload, load_index, ChunkRecord
from app.services.runtime_config import get_runtime_overrides_sync, get_llm_config
from app.services.llm import generate_chat_completion
from app.utils.responses import raise_http_error, with_corr_id


def get_collection() -> AsyncIOMotorCollection:
    return mongo_collection("rags")


async def ensure_indexes() -> None:
    col = get_collection()
    await col.create_index("slug", unique=True)
    await col.create_index("users")
    await col.create_index("docs.doc_id")


ingest_logger = logging.getLogger("chatfleet.rag.ingest")
ingest_logger.setLevel(logging.INFO)

ALLOWED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".odt": "application/vnd.oasis.opendocument.text",
}


def _doc_to_summary(doc: Dict[str, Any]) -> RagSummary:
    return RagSummary(
        slug=doc["slug"],
        name=doc.get("name", doc["slug"]),
        description=doc.get("description", ""),
        chunks=doc.get("chunks", 0),
        last_updated=doc.get("last_updated", datetime.now(timezone.utc)),
        visibility=doc.get("visibility", "private"),
        suggestions=doc.get("suggestions", []) or [],
        suggestions_en=doc.get("suggestions_en", []) or [],
        suggestions_lang=doc.get("suggestions_lang"),
    )


def _doc_to_rag_doc(entry: Dict[str, Any]) -> RagDoc:
    return RagDoc(
        doc_id=entry["doc_id"],
        filename=entry["filename"],
        path=entry.get("path"),
        mime=entry.get("mime", "application/pdf"),
        size_bytes=entry.get("size_bytes", 0),
        sha256=entry.get("sha256"),
        status=entry.get("status", "uploaded"),
        chunk_count=entry.get("chunk_count", 0),
        error=entry.get("error"),
        uploaded_at=entry.get("uploaded_at"),
        indexed_at=entry.get("indexed_at"),
    )


async def _extract_pdf_text(path: str) -> List[PageText]:
    def _normalize(text: str) -> str:
        import re

        # Normalize newlines, collapse repeated whitespace, keep paragraph breaks.
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"-\n(?=[a-z])", "", text)
        # undo hyphenation across lines
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _read_pypdf() -> List[PageText]:
        reader = PdfReader(path)
        pages: List[PageText] = []
        for idx, page in enumerate(reader.pages):
            try:
                snippet = page.extract_text() or ""
            except Exception:
                snippet = ""
            normalized = _normalize(snippet)
            if normalized:
                pages.append(PageText(page=idx + 1, text=normalized))
        return pages

    def _read_pdfplumber() -> List[PageText]:
        try:
            import pdfplumber
        except Exception:
            return []

        pages: List[PageText] = []
        with pdfplumber.open(path) as pdf:
            for idx, page in enumerate(pdf.pages):
                snippet = page.extract_text(layout=True) or page.extract_text() or ""
                normalized = _normalize(snippet)
                if normalized:
                    pages.append(PageText(page=idx + 1, text=normalized))
        return pages

    def _read() -> List[PageText]:
        # Prefer layout-aware extraction; fall back to PyPDF2 if unavailable.
        pages = _read_pdfplumber()
        return pages or _read_pypdf()

    pages = await asyncio.to_thread(_read)
    if not pages:
        raise ValueError("No extractable text found in PDF")
    return pages


async def _extract_docx_text(path: str) -> List[PageText]:
    def _read() -> List[PageText]:
        doc = DocxDocument(path)
        pages: List[PageText] = []
        page = 1
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                pages.append(PageText(page=page, text=text))
                page += 1
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    pages.append(PageText(page=page, text=row_text))
                    page += 1
        return pages

    pages = await asyncio.to_thread(_read)
    if not pages:
        raise ValueError("No extractable text found in DOCX")
    return pages


async def _extract_odt_text(path: str) -> List[PageText]:
    def _read() -> List[PageText]:
        doc = load_odt(path)
        pages: List[PageText] = []
        page = 1

        for para in doc.getElementsByType(OdtParagraph):
            text = teletype.extractText(para).strip()
            if text:
                pages.append(PageText(page=page, text=text))
                page += 1

        for table in doc.getElementsByType(OdtTable):
            for row in table.getElementsByType(OdtTableRow):
                cells = [
                    teletype.extractText(cell).strip()
                    for cell in row.getElementsByType(OdtTableCell)
                    if teletype.extractText(cell).strip()
                ]
                if cells:
                    pages.append(PageText(page=page, text=" | ".join(cells)))
                    page += 1

        return pages

    pages = await asyncio.to_thread(_read)
    if not pages:
        raise ValueError("No extractable text found in ODT")
    return pages


async def _extract_txt_text(path: str) -> List[PageText]:
    def _read() -> List[PageText]:
        raw = Path(path).read_bytes()
        detected = chardet.detect(raw) if raw else {}
        encoding = (detected.get("encoding") or "utf-8").lower()
        text = raw.decode(encoding, errors="replace")
        text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return []
        return [PageText(page=1, text=text)]

    pages = await asyncio.to_thread(_read)
    if not pages:
        raise ValueError("No extractable text found in text file")
    return pages


async def _persist_chunks_and_vectors(
    rag_slug: str,
    doc_entry: Dict[str, Any],
    chunks: Sequence[ChunkWithPage],
) -> int:
    if not chunks:
        raise ValueError("No textual chunks extracted")
    embeddings = await embed_texts([chunk.text for chunk in chunks])
    if len(embeddings) != len(chunks):
        raise ValueError("Embedding count does not match chunk count")
    # Enforce consistent embedding dimension per RAG
    dim = len(embeddings[0]) if embeddings and embeddings[0] is not None else 0
    if embeddings:
        lengths = [len(chunk.text) for chunk in chunks]
        ingest_logger.info(
            "rag.ingest.embed rag=%s doc=%s file=%s chunks=%s dim=%s chars[min/avg/max]=[%s/%.1f/%s]",
            rag_slug,
            doc_entry["doc_id"],
            doc_entry["filename"],
            len(chunks),
            dim,
            min(lengths) if lengths else 0,
            (sum(lengths) / len(lengths)) if lengths else 0.0,
            max(lengths) if lengths else 0,
            extra={
                "rag_slug": rag_slug,
                "doc_id": doc_entry["doc_id"],
                "file_name": doc_entry["filename"],
                "chunks": len(chunks),
                "dim": dim,
                "min_chars": min(lengths) if lengths else 0,
                "max_chars": max(lengths) if lengths else 0,
                "avg_chars": sum(lengths) / len(lengths) if lengths else 0,
            },
        )
    try:
        current_rag = await get_rag_by_slug(rag_slug)
    except Exception:
        current_rag = None
    if current_rag:
        # Prefer existing index dim if built
        existing_dim = 0
        idx = current_rag.get("index", {}) if isinstance(current_rag, dict) else {}
        if isinstance(idx, dict):
            existing_dim = int(idx.get("dim") or 0)
        embed_dim_field = int(current_rag.get("embed_dim") or 0)
        guard_dim = existing_dim or embed_dim_field
        if guard_dim and dim and guard_dim != dim:
            raise ValueError(f"EMBED_DIM_MISMATCH: current={dim} existing={guard_dim}")
        # Persist embed_dim on first embed if not set
        if not guard_dim and dim:
            col = get_collection()
            await col.update_one({"slug": rag_slug}, {"$set": {"embed_dim": int(dim)}})
    persist_doc_payload(rag_slug, doc_entry["doc_id"], doc_entry["filename"], chunks, embeddings)
    return len(chunks)


async def _update_rag_index_summary(rag_slug: str, total_chunks: int, dimension: int) -> None:
    col = get_collection()
    await col.update_one(
        {"slug": rag_slug},
        {
            "$set": {
                "chunks": total_chunks,
                "index": {
                    "type": "faiss",
                    "path": str(settings.index_dir / rag_slug / "index.faiss"),
                    # Store runtime embed model name for observability
                    "emb_model": (await get_llm_config()).embed_model if dimension else settings.embed_model,
                    "dim": dimension,
                    "built_at": datetime.now(timezone.utc),
                },
                "last_updated": datetime.now(timezone.utc),
            }
        },
    )


def _select_chunk_samples(records: List[ChunkRecord], limit: int = 32, max_chars: int = 320) -> List[str]:
    if not records:
        return []
    step = max(1, len(records) // max(1, min(limit, len(records))))
    sampled = records[::step][:limit]
    seen: set[str] = set()
    snippets: List[str] = []
    for rec in sampled:
        text = (rec.text or "").strip()
        if not text:
            continue
        snippet = text[:max_chars].strip()
        normalized = snippet.lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            snippets.append(snippet)
    return snippets


def _normalize_suggestions_list(candidates: List[str], limit: int = 6, max_len: int = 180) -> List[str]:
    out: List[str] = []

    def push(val: str) -> None:
        cleaned = val.strip().strip('"').strip()
        if not cleaned:
            return
        # If a single string contains multiple prompts, split on common separators.
        if cleaned.count('","') or cleaned.count('\" , \"') or cleaned.count('", "'):
            parts = re.split(r'"\s*,\s*"', cleaned)
            for part in parts:
                push(part)
            return
        if cleaned.startswith("[") and cleaned.endswith("]"):
            try:
                parsed = json.loads(cleaned)
                if isinstance(parsed, list):
                    for item in parsed:
                        if isinstance(item, str):
                            push(item)
                    return
            except Exception:
                pass
        cleaned = cleaned[:max_len].strip()
        if cleaned and cleaned.lower() not in {x.lower() for x in out}:
            out.append(cleaned)

    for item in candidates:
        push(str(item))

    return out[:limit]


async def _load_chunk_samples_for_suggestions(rag_slug: str, limit: int = 32, max_chars: int = 320) -> List[str]:
    try:
        _, records = load_index(rag_slug)
    except FileNotFoundError:
        return []
    except Exception as exc:
        ingest_logger.warning(
            "rag.suggestions.sample_failed rag=%s error=%s", rag_slug, exc, extra={"rag_slug": rag_slug}
        )
        return []
    return _select_chunk_samples(records, limit=limit, max_chars=max_chars)


def _parse_suggestions(raw: str, limit: int = 4, max_len: int = 120) -> List[str]:
    raw = (raw or "").strip()
    if not raw:
        return []

    def _clean(item: str) -> str:
        item = re.sub(r"^\s*[-*\d\.\)]\s*", "", item.strip())
        item = item.strip().strip('"').strip("'")
        return item[:max_len].strip()

    suggestions: List[str] = []
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            for entry in parsed:
                if isinstance(entry, str):
                    cleaned = _clean(entry)
                    if cleaned:
                        suggestions.append(cleaned)
                elif isinstance(entry, dict):
                    text = entry.get("q") or entry.get("question") or ""
                    cleaned = _clean(str(text))
                    if cleaned:
                        suggestions.append(cleaned)
    except Exception:
        pass

    if not suggestions:
        for line in raw.splitlines():
            cleaned = _clean(line)
            if cleaned:
                suggestions.append(cleaned)

    deduped: List[str] = []
    seen: set[str] = set()
    for entry in suggestions:
        key = entry.lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(entry)
        if len(deduped) >= limit:
            break
    return deduped


def _detect_language(snippets: List[str], description: str) -> str:
    sample_text = (" ".join(snippets[:10]) or description or "").lower()
    french_markers = [
        "le ",
        "la ",
        "les ",
        "des ",
        "un ",
        "une ",
        "pour ",
        "comment ",
        "quelles ",
        "quels ",
        "guide ",
        "manuel ",
        "chapitre ",
    ]
    has_french_char = bool(re.search(r"[àâçéèêëîïôûùüÿœ]", sample_text))
    has_marker = any(marker in sample_text for marker in french_markers)
    return "fr" if has_french_char or has_marker else "en"


def _filter_by_overlap(candidates: List[str], snippets: List[str]) -> List[str]:
    if not snippets:
        return candidates
    kept: List[str] = []
    snippet_tokens = [set(re.findall(r"[A-Za-zÀ-ÿ0-9]{3,}", s.lower())) for s in snippets]
    for q in candidates:
        q_tokens = set(re.findall(r"[A-Za-zÀ-ÿ0-9]{3,}", q.lower()))
        match = any(len(q_tokens & tokens) >= 2 for tokens in snippet_tokens)
        if match:
            kept.append(q)
    return kept or candidates


async def _generate_rag_suggestions(
    rag_slug: str, name: str, description: str
) -> tuple[List[str], List[str], str, Optional[str]]:
    snippets = await _load_chunk_samples_for_suggestions(rag_slug, limit=32, max_chars=320)
    lang_primary = _detect_language(snippets, description)

    payload_summary = f"Assistant name: {name or rag_slug}\nDescription: {description or 'N/A'}"
    snippet_block = "\n".join(f"{idx+1}) {s}" for idx, s in enumerate(snippets[:32]))
    ingest_logger.info(
        "rag.suggestions.prompt rag=%s snippets=%s len=%s lang=%s",
        rag_slug,
        len(snippets),
        len(snippet_block),
        lang_primary,
        extra={"rag_slug": rag_slug, "snippets": len(snippets), "lang": lang_primary},
    )

    base_prompt = [
        {
            "role": "system",
            "content": (
                "You generate 3-4 concise end-user questions for a documentation assistant. "
                "Each question MUST be answerable directly from the provided snippets; avoid topics not covered. "
                "Keep each question under 120 characters. "
                "Return ONLY a JSON array; each element should be an object: "
                "{\"q\": \"question text\", \"snippets\": [numbers of supporting snippets]}. "
                "No markdown, no numbering. Language: {{lang}}."
            ),
        },
        {
            "role": "user",
            "content": f"{payload_summary}\n\nSnippets:\n{snippet_block}",
        },
    ]

    async def _run_for_lang(lang: str) -> List[str]:
        messages = [
            {**base_prompt[0], "content": base_prompt[0]["content"].replace("{{lang}}", "French" if lang == "fr" else "English")},
            base_prompt[1],
        ]
        llm_result = await generate_chat_completion(messages, temperature=0.3, max_tokens=320)
        suggestions: List[str] = []
        if llm_result is not None:
            raw, _ = llm_result
            ingest_logger.info(
                "rag.suggestions.raw rag=%s lang=%s text=%s",
                rag_slug,
                lang,
                raw[:500],
                extra={"rag_slug": rag_slug, "lang": lang},
            )
            suggestions = _parse_suggestions(raw)

        if not suggestions:
            base = name or rag_slug
            if lang == "fr":
                suggestions = [
                    f"Peux-tu résumer rapidement {base} ?",
                    f"Comment démarrer avec {base} ?",
                    "Quelles sont les notions clés à retenir ?",
                    "Donne-moi un exemple pratique tiré de cette documentation.",
                ]
            else:
                suggestions = [
                    f"Can you give a quick overview of {base}?",
                    f"How do I get started with {base}?",
                    "What are the key concepts covered?",
                    "Share a practical example from this documentation.",
                ]
        normalized = _normalize_suggestions_list(
            [s.strip(" []\"'") for s in suggestions if s.strip(" []\"'")],
            limit=4,
        )
        normalized = _filter_by_overlap(normalized, snippets)
        return normalized

    primary = await _run_for_lang(lang_primary)
    secondary_lang = "en" if lang_primary == "fr" else "fr"
    secondary = await _run_for_lang(secondary_lang)

    ingest_logger.info(
        "rag.suggestions.final rag=%s lang_primary=%s primary=%s secondary_lang=%s secondary=%s",
        rag_slug,
        lang_primary,
        len(primary),
        secondary_lang,
        len(secondary),
        extra={
            "rag_slug": rag_slug,
            "lang_primary": lang_primary,
            "primary": primary,
            "secondary_lang": secondary_lang,
            "secondary": secondary,
        },
    )
    return primary, secondary, lang_primary, None


async def _persist_rag_suggestions(
    rag_slug: str,
    primary: List[str],
    secondary: List[str],
    lang_primary: str,
) -> None:
    normalized_primary = _normalize_suggestions_list(primary, limit=6)
    normalized_secondary = _normalize_suggestions_list(secondary, limit=6)
    # Ensure we always store English suggestions when available
    suggestions_en = normalized_secondary if lang_primary == "fr" else normalized_primary
    ingest_logger.info(
        "rag.suggestions.store rag=%s lang_primary=%s primary=%s secondary=%s",
        rag_slug,
        lang_primary,
        normalized_primary,
        normalized_secondary,
        extra={
            "rag_slug": rag_slug,
            "lang_primary": lang_primary,
            "primary": normalized_primary,
            "secondary": normalized_secondary,
        },
    )
    col = get_collection()
    await col.update_one(
        {"slug": rag_slug},
        {
            "$set": {
                "suggestions": normalized_primary,
                "suggestions_en": suggestions_en,
                "suggestions_lang": lang_primary,
                "suggestions_updated_at": datetime.now(timezone.utc),
                "last_updated": datetime.now(timezone.utc),
            }
        },
    )


async def list_rags_for_slugs(slugs: List[str], limit: int = 50, cursor: Optional[str] = None, include_public: bool = False) -> Tuple[List[RagSummary], Optional[str]]:
    col = get_collection()
    ors: List[Dict[str, Any]] = []
    if slugs:
        ors.append({"slug": {"$in": slugs}})
    if include_public:
        ors.append({"visibility": "public"})
    if not ors:
        return [], None
    query: Dict[str, Any] = {"$or": ors}
    if cursor:
        query["_id"] = {"$gt": ObjectId(cursor)}
    docs: List[Dict[str, Any]] = []
    async for doc in col.find(query).sort("_id", 1).limit(limit + 1):
        docs.append(doc)
    next_cursor = None
    if len(docs) > limit:
        next_cursor = str(docs[-1]["_id"])
        docs = docs[:-1]
    return [_doc_to_summary(doc) for doc in docs], next_cursor


async def list_all_rags(limit: int = 50, cursor: Optional[str] = None) -> Tuple[List[RagSummary], Optional[str]]:
    col = get_collection()
    query: Dict[str, Any] = {}
    if cursor:
        query["_id"] = {"$gt": ObjectId(cursor)}
    docs: List[Dict[str, Any]] = []
    async for doc in col.find(query).sort("_id", 1).limit(limit + 1):
        docs.append(doc)
    next_cursor = None
    if len(docs) > limit:
        next_cursor = str(docs[-1]["_id"])
        docs = docs[:-1]
    return [_doc_to_summary(doc) for doc in docs], next_cursor


async def list_public_rags(limit: int = 50, cursor: Optional[str] = None) -> Tuple[List[RagSummary], Optional[str]]:
    col = get_collection()
    query: Dict[str, Any] = {"visibility": "public"}
    if cursor:
        query["_id"] = {"$gt": ObjectId(cursor)}
    docs: List[Dict[str, Any]] = []
    async for doc in col.find(query).sort("_id", 1).limit(limit + 1):
        docs.append(doc)
    next_cursor = None
    if len(docs) > limit:
        next_cursor = str(docs[-1]["_id"])
        docs = docs[:-1]
    return [_doc_to_summary(doc) for doc in docs], next_cursor


async def get_rag_by_slug(slug: str) -> Optional[Dict[str, Any]]:
    col = get_collection()
    return await col.find_one({"slug": slug})


async def create_rag(payload: RagCreateRequest, creator_id: str) -> RagSummary:
    col = get_collection()
    now = datetime.now(timezone.utc)
    doc: Dict[str, Any] = {
        "slug": payload.slug,
        "name": payload.name,
        "description": payload.description,
        "users": [],
        "docs": [],
        "chunks": 0,
        "visibility": payload.visibility or "private",
        "suggestions": [],
        "index": {
            "type": "faiss",
            "path": "",
            "emb_model": settings.embed_model,
            "dim": 0,
            "built_at": None,
        },
        "created_at": now,
        "last_updated": now,
        "created_by": ObjectId(creator_id),
    }
    try:
        result = await col.insert_one(doc)
    except DuplicateKeyError:
        raise_http_error("RAG_EXISTS", f"RAG '{payload.slug}' already exists", status.HTTP_409_CONFLICT)
    doc["_id"] = result.inserted_id
    await write_system_log(
        event="rag.create",
        rag_slug=payload.slug,
        user_id=creator_id,
        details={"name": payload.name},
    )
    return _doc_to_summary(doc)


async def get_docs(slug: str) -> RagDocsResponse:
    rag = await get_rag_by_slug(slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{slug}' not found", status_code=404)
    docs = [_doc_to_rag_doc(entry) for entry in rag.get("docs", [])]
    return RagDocsResponse(**with_corr_id({"rag_slug": slug, "docs": [doc.model_dump() for doc in docs]}))


async def get_index_status(slug: str) -> IndexStatusResponse:
    rag = await get_rag_by_slug(slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{slug}' not found", status_code=404)
    docs = rag.get("docs", [])
    total_docs = len(docs)
    done_docs = sum(1 for doc in docs if doc.get("status") == "indexed")
    total_chunks = sum(doc.get("chunk_count", 0) for doc in docs)
    done_chunks = total_chunks if done_docs == total_docs else sum(
        doc.get("chunk_count", 0) for doc in docs if doc.get("status") == "indexed"
    )
    status: IndexStatusEnum = "idle"
    if any(doc.get("status") in {"chunking", "indexing"} for doc in docs):
        status = "building"
    if any(doc.get("status") == "error" for doc in docs):
        status = "error"
    progress = 0.0 if total_docs == 0 else done_docs / total_docs
    index_info = rag.get("index", {})
    error = index_info.get("error")
    payload = {
        "rag_slug": slug,
        "status": status,
        "progress": progress,
        "total_docs": total_docs,
        "done_docs": done_docs,
        "total_chunks": total_chunks,
        "done_chunks": done_chunks,
        "error": error,
    }
    return IndexStatusResponse(**with_corr_id(payload))


def _save_upload_target(rag_slug: str) -> Path:
    _, upload_dir, _, _, _ = get_runtime_overrides_sync()
    target_dir = upload_dir / rag_slug
    target_dir.mkdir(parents=True, exist_ok=True)
    return target_dir


async def _persist_file(rag_slug: str, upload: UploadFile, doc_id: str, extension: str) -> Tuple[str, int, str]:
    """Persist an uploaded file safely.

    - Store using a generated filename `<doc_id><extension>` to avoid path traversal.
    - Validate basic PDF magic header for PDF uploads.
    - Enforce max size during streaming to avoid excessive disk usage.
    """

    target_dir = _save_upload_target(rag_slug)
    safe_filename = f"{doc_id}{extension}"
    target_path = target_dir / safe_filename
    size = 0
    hasher = hashlib.sha256()
    max_bytes = settings.max_upload_mb * 1024 * 1024

    first_chunk = await upload.read(8 * 1024)
    if not first_chunk:
        await upload.close()
        raise ValueError("Empty upload")
    # Basic PDF signature check
    if extension == ".pdf" and not first_chunk.startswith(b"%PDF"):
        await upload.close()
        raise ValueError("Only PDF files are accepted")

    with target_path.open("wb") as out_file:
        size += len(first_chunk)
        if size > max_bytes:
            await upload.close()
            out_file.close()
            try:
                target_path.unlink()
            except FileNotFoundError:
                pass
            raise ValueError(f"{upload.filename or safe_filename} exceeds {settings.max_upload_mb}MB limit")
        hasher.update(first_chunk)
        out_file.write(first_chunk)

        while True:
            chunk = await upload.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes:
                await upload.close()
                out_file.close()
                try:
                    target_path.unlink()
                except FileNotFoundError:
                    pass
                raise ValueError(f"{upload.filename or safe_filename} exceeds {settings.max_upload_mb}MB limit")
            hasher.update(chunk)
            out_file.write(chunk)
    await upload.close()
    return str(target_path), size, hasher.hexdigest()


async def _register_doc_metadata(
    rag: Dict[str, Any],
    filename: str,
    mime: str,
    stored_path: str,
    size_bytes: int,
    sha256: str,
) -> Dict[str, Any]:
    doc_id = str(uuid4())
    entry = {
        "doc_id": doc_id,
        "filename": filename,
        "path": stored_path,
        "mime": mime,
        "size_bytes": size_bytes,
        "sha256": sha256,
        "status": "uploaded",
        "chunk_count": 0,
        "error": None,
        "uploaded_at": datetime.now(timezone.utc),
        "indexed_at": None,
    }
    docs = rag.get("docs", [])
    docs.append(entry)
    rag["docs"] = docs
    rag["last_updated"] = datetime.now(timezone.utc)
    col = get_collection()
    await col.update_one({"_id": rag["_id"]}, {"$set": {"docs": docs, "last_updated": rag["last_updated"]}})
    return entry


async def upload_documents(
    rag_slug: str,
    uploads: List[UploadFile],
    uploader_id: str,
) -> RagUploadAccepted:
    rag = await get_rag_by_slug(rag_slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{rag_slug}' not found", status_code=404)

    accepted: List[str] = []
    skipped: List[str] = []
    new_entries: List[Dict[str, Any]] = []

    for upload in uploads:
        ext = Path(upload.filename or "").suffix.lower()
        mime = ALLOWED_EXTENSIONS.get(ext)
        if not mime:
            skipped.append(upload.filename or "unnamed")
            continue
        try:
            stored_path, size_bytes, sha256 = await _persist_file(rag_slug, upload, str(uuid4()), ext)
        except ValueError as exc:
            skipped.append(upload.filename or "unnamed")
            await write_system_log(
                event="rag.upload.skipped",
                rag_slug=rag_slug,
                user_id=uploader_id,
                details={"reason": str(exc)},
                level="warn",
            )
            continue
        entry = await _register_doc_metadata(
            rag,
            upload.filename or Path(stored_path).name,
            mime,
            stored_path,
            size_bytes,
            sha256,
        )
        new_entries.append(entry)
        accepted.append(entry["filename"])

    async def runner(job: JobRecord) -> None:
        chunk_counts: Dict[str, int] = {}
        totals = {
            "docs_total": len(new_entries),
            "docs_done": 0,
            "chunks_total": 0,
            "chunks_done": 0,
        }

        def _apply_progress(phase: Optional[str] = None) -> None:
            job.phase = phase or job.phase
            job.totals = JobProgressTotals(**totals)
            job.progress = _compute_progress(totals)

        def _compute_progress(payload: Dict[str, int]) -> float:
            docs_total = payload.get("docs_total", 0)
            docs_done = payload.get("docs_done", 0)
            chunks_total = payload.get("chunks_total", 0)
            chunks_done = payload.get("chunks_done", 0)
            if chunks_total > 0:
                return min(0.98, chunks_done / max(chunks_total, 1))
            if docs_total > 0:
                return min(0.98, docs_done / max(docs_total, 1))
            return 1.0

        try:
            if not new_entries:
                await write_system_log(
                    event="rag.upload.noop",
                    rag_slug=rag_slug,
                    user_id=uploader_id,
                    details={"job_id": job.job_id, "skipped": skipped},
                    level="warn",
                )
                job.progress = 1.0
                job.phase = "finalizing"
                job.totals = JobProgressTotals(
                    docs_total=0, docs_done=0, chunks_total=0, chunks_done=0
                )
                job.suggestions_ready = True
                job.result = {"total_chunks": 0, "dimension": 0, "suggestions": 0}
                return
            _apply_progress("queued")
            for entry in new_entries:
                doc_id = entry["doc_id"]
                try:
                    await _update_doc_status(rag_slug, doc_id, "chunking")
                    _apply_progress("chunking")
                    mime = entry.get("mime", "application/pdf")
                    if mime == "application/pdf":
                        pages = await _extract_pdf_text(entry["path"])
                    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        pages = await _extract_docx_text(entry["path"])
                    elif mime == "application/vnd.oasis.opendocument.text":
                        pages = await _extract_odt_text(entry["path"])
                    elif mime == "text/plain":
                        pages = await _extract_txt_text(entry["path"])
                    elif mime == "application/msword":
                        raise ValueError("DOC files are not supported; convert to DOCX")
                    else:
                        raise ValueError(f"Unsupported mime type: {mime}")
                    chunks = chunk_pdf(pages)
                    chunk_count = len(chunks)
                    totals["chunks_total"] += chunk_count
                    _apply_progress("chunking")
                    chunk_counts[doc_id] = chunk_count
                    ingest_logger.info(
                        "rag.ingest.chunk rag=%s doc=%s file=%s pages=%s chunks=%s",
                        rag_slug,
                        doc_id,
                        entry["filename"],
                        len(pages),
                        chunk_count,
                        extra={
                            "rag_slug": rag_slug,
                            "doc_id": doc_id,
                            "file_name": entry["filename"],
                            "pages": len(pages),
                            "chunks": chunk_count,
                            "preview": [
                                {"page_start": chunk.page_start, "page_end": chunk.page_end, "chars": len(chunk.text)}
                                for chunk in chunks[:5]
                            ],
                        },
                    )
                    await _update_doc_status(rag_slug, doc_id, "chunked", chunk_count=chunk_count)
                    _apply_progress("embedding")
                    stored_count = await _persist_chunks_and_vectors(rag_slug, entry, chunks)
                    totals["chunks_done"] += stored_count
                    totals["docs_done"] += 1
                    await _update_doc_status(rag_slug, doc_id, "indexing", chunk_count=stored_count)
                    _apply_progress("indexing")
                except Exception as doc_exc:
                    await _update_doc_status(rag_slug, doc_id, "error", error=str(doc_exc))
                    raise

            job.phase = "indexing"
            total_chunks, dimension = await asyncio.to_thread(rebuild_vector_index, rag_slug)
            for entry in new_entries:
                doc_id = entry["doc_id"]
                await _update_doc_status(rag_slug, doc_id, "indexed", chunk_count=chunk_counts.get(doc_id, 0))
            totals["chunks_total"] = max(totals["chunks_total"], total_chunks)
            totals["chunks_done"] = max(totals["chunks_done"], total_chunks)
            _apply_progress("finalizing")
            await _update_rag_index_summary(rag_slug, total_chunks, dimension)

            job.phase = "suggestions"
            suggestions: List[str] = []
            suggestions_secondary: List[str] = []
            suggestions_error: Optional[str] = None
            lang_primary = "fr"
            try:
                suggestions, suggestions_secondary, lang_primary, suggestions_error = await _generate_rag_suggestions(
                    rag_slug,
                    rag.get("name", rag_slug),
                    rag.get("description", ""),
                )
            except Exception as exc:
                suggestions_error = str(exc)
                ingest_logger.warning(
                    "rag.suggestions.generate_failed rag=%s error=%s", rag_slug, exc, extra={"rag_slug": rag_slug}
                )
            finally:
                await _persist_rag_suggestions(
                    rag_slug,
                    suggestions,
                    suggestions_secondary,
                    lang_primary,
                )
                job.suggestions_ready = True
                job.result = {
                    "total_chunks": total_chunks,
                    "dimension": dimension,
                    "suggestions": len(suggestions),
                    "suggestions_secondary": len(suggestions_secondary),
                    "suggestions_lang": lang_primary,
                }
                if suggestions_error:
                    job.result["suggestions_error"] = suggestions_error
            job.progress = 1.0
            await write_system_log(
                event="rag.upload.complete",
                rag_slug=rag_slug,
                user_id=uploader_id,
                details={
                    "accepted": accepted,
                    "skipped": skipped,
                    "job_id": job.job_id,
                    "chunks": total_chunks,
                    "suggestions": len(suggestions),
                    "suggestions_error": suggestions_error,
                },
            )
        except Exception as exc:  # pragma: no cover — safety net
            job.status = "error"
            job.error = str(exc)
            await write_system_log(
                event="rag.upload.failed",
                rag_slug=rag_slug,
                user_id=uploader_id,
                details={"error": str(exc), "job_id": job.job_id},
                level="error",
            )
            raise

    job = job_manager.schedule("RAG_INDEX", runner)
    payload = {
        "job_id": job.job_id,
        "accepted": accepted,
        "skipped": skipped,
        "rag_slug": rag_slug,
    }
    return RagUploadAccepted(**with_corr_id(payload))


async def _update_doc_status(
    rag_slug: str,
    doc_id: str,
    status: str,
    chunk_count: Optional[int] = None,
    error: Optional[str] = None,
) -> None:
    col = get_collection()
    await col.update_one(
        {"slug": rag_slug, "docs.doc_id": doc_id},
        {
            "$set": {
                "docs.$.status": status,
                "docs.$.chunk_count": chunk_count if chunk_count is not None else 0,
                "docs.$.indexed_at": datetime.now(timezone.utc) if status == "indexed" else None,
                "docs.$.error": error,
                "last_updated": datetime.now(timezone.utc),
            }
        },
    )


async def rebuild_index(rag_slug: str, user_id: str) -> str:
    rag = await get_rag_by_slug(rag_slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{rag_slug}' not found", status_code=404)

    async def runner(job: JobRecord) -> None:
        try:
            current = await get_rag_by_slug(rag_slug)
            docs = current.get("docs", []) if current else []
            for doc in docs:
                if doc.get("status") != "error":
                    await _update_doc_status(
                        rag_slug,
                        doc["doc_id"],
                        "indexing",
                        chunk_count=doc.get("chunk_count", 0),
                    )

            total_chunks, dimension = await asyncio.to_thread(rebuild_vector_index, rag_slug)

            current = await get_rag_by_slug(rag_slug)
            docs = current.get("docs", []) if current else []
            for doc in docs:
                if doc.get("status") != "error":
                    await _update_doc_status(
                        rag_slug,
                        doc["doc_id"],
                        "indexed",
                        chunk_count=doc.get("chunk_count", 0),
                    )

            await _update_rag_index_summary(rag_slug, total_chunks, dimension)

            job.phase = "suggestions"
            suggestions: List[str] = []
            suggestions_secondary: List[str] = []
            suggestions_error: Optional[str] = None
            lang_primary = "fr"
            try:
                suggestions, suggestions_secondary, lang_primary, suggestions_error = await _generate_rag_suggestions(
                    rag_slug,
                    rag.get("name", rag_slug),
                    rag.get("description", ""),
                )
            except Exception as exc:
                suggestions_error = str(exc)
                ingest_logger.warning(
                    "rag.suggestions.generate_failed rag=%s error=%s", rag_slug, exc, extra={"rag_slug": rag_slug}
                )
            finally:
                await _persist_rag_suggestions(
                    rag_slug,
                    suggestions,
                    suggestions_secondary,
                    lang_primary,
                )
                job.suggestions_ready = True
                job.result = {
                    "total_chunks": total_chunks,
                    "dimension": dimension,
                    "suggestions": len(suggestions),
                    "suggestions_secondary": len(suggestions_secondary),
                    "suggestions_lang": lang_primary,
                    "suggestions_error": suggestions_error,
                }

            await write_system_log(
                event="rag.rebuild.complete",
                rag_slug=rag_slug,
                user_id=user_id,
                details={
                    "job_id": job.job_id,
                    "chunks": total_chunks,
                    "suggestions": len(suggestions),
                    "suggestions_error": suggestions_error,
                },
            )
        except Exception as exc:  # pragma: no cover — defensive
            job.status = "error"
            job.error = str(exc)
            await write_system_log(
                event="rag.rebuild.failed",
                rag_slug=rag_slug,
                user_id=user_id,
                details={"error": str(exc), "job_id": job.job_id},
                level="error",
            )
            raise

    job = job_manager.schedule("RAG_REBUILD", runner)
    return job.job_id


async def reset_index(rag_slug: str, user_id: str) -> str:
    rag = await get_rag_by_slug(rag_slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{rag_slug}' not found", status_code=404)

    async def runner(job: JobRecord) -> None:
        col = get_collection()
        await col.update_one(
            {"_id": rag["_id"]},
            {
                "$set": {
                    "docs": [],
                    "chunks": 0,
                    "index": {
                        "type": "faiss",
                        "path": "",
                        "emb_model": (await get_llm_config()).embed_model,
                        "dim": 0,
                        "built_at": None,
                    },
                    "embed_dim": 0,
                    "last_updated": datetime.now(timezone.utc),
                }
            },
        )

        for doc in rag.get("docs", []):
            remove_doc_payload(rag_slug, doc["doc_id"])

        index_dir, upload_dir_base, _, _, _ = get_runtime_overrides_sync()
        index_dir = index_dir / rag_slug
        if index_dir.exists():
            for child in index_dir.glob("*"):
                try:
                    if child.is_file():
                        child.unlink()
                    elif child.is_dir():
                        for nested in child.rglob("*"):
                            if nested.is_file():
                                nested.unlink()
                        child.rmdir()
                except OSError:
                    pass

        target_dir = upload_dir_base / rag_slug
        if target_dir.exists():
            for child in target_dir.glob("*"):
                try:
                    if child.is_file():
                        child.unlink()
                    elif child.is_dir():
                        for nested in child.rglob("*"):
                            if nested.is_file():
                                nested.unlink()
                        child.rmdir()
                except OSError:
                    pass

        job.result = {"total_chunks": 0}
        await write_system_log(
            event="rag.reset.complete",
            rag_slug=rag_slug,
            user_id=user_id,
            details={"job_id": job.job_id},
        )

    job = job_manager.schedule("RAG_RESET", runner)
    return job.job_id


async def list_rag_users(rag_slug: str) -> RagUsersResponse:
    rag = await get_rag_by_slug(rag_slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{rag_slug}' not found", status_code=404)
    user_ids: List[ObjectId] = rag.get("users", [])
    users: List[RagUser] = []
    for user_id in user_ids:
        doc = await find_user_by_id(str(user_id))
        if doc:
            users.append(RagUser(id=str(doc["_id"]), email=doc["email"], name=doc["name"], role=doc.get("role", "user")))
    payload = {"rag_slug": rag_slug, "users": [user.model_dump(by_alias=True) for user in users]}
    return RagUsersResponse(**with_corr_id(payload))


async def add_user_access(rag_slug: str, user_id: ObjectId) -> None:
    col = get_collection()
    await col.update_one({"slug": rag_slug}, {"$addToSet": {"users": user_id}})


async def remove_user_access(rag_slug: str, user_id: ObjectId) -> None:
    col = get_collection()
    await col.update_one({"slug": rag_slug}, {"$pull": {"users": user_id}})


async def add_user_to_rag(rag_slug: str, user_id: ObjectId, current_rags: List[str]) -> None:
    current = set(current_rags)
    current.add(rag_slug)
    await update_user_rags(user_id, sorted(current))
    await add_user_access(rag_slug, user_id)


async def remove_user_from_rag(rag_slug: str, user_id: ObjectId, current_rags: List[str]) -> None:
    current = [slug for slug in current_rags if slug != rag_slug]
    await update_user_rags(user_id, current)
    await remove_user_access(rag_slug, user_id)


async def delete_rag(rag_slug: str, confirmation: str, actor_id: str) -> None:
    if confirmation != rag_slug:
        raise_http_error("CONFIRMATION_MISMATCH", "Confirmation must match rag_slug", status.HTTP_422_UNPROCESSABLE_ENTITY)

    rag = await get_rag_by_slug(rag_slug)
    if not rag:
        raise_http_error("RAG_NOT_FOUND", f"RAG '{rag_slug}' not found", status.HTTP_404_NOT_FOUND)

    # Optionally check for lock or concurrent operations here (placeholder for future state).

    # Remove document payloads and indexes.
    for doc in rag.get("docs", []):
        remove_doc_payload(rag_slug, doc["doc_id"])

    index_dir, upload_dir_base, _, _, _ = get_runtime_overrides_sync()
    index_dir = index_dir / rag_slug
    if index_dir.exists():
        for child in index_dir.glob("*"):
            try:
                if child.is_file():
                    child.unlink()
                elif child.is_dir():
                    for nested in child.rglob("*"):
                        if nested.is_file():
                            nested.unlink()
                    child.rmdir()
            except OSError:
                pass
        try:
            index_dir.rmdir()
        except OSError:
            pass

    upload_dir = upload_dir_base / rag_slug
    if upload_dir.exists():
        for child in upload_dir.glob("*"):
            try:
                if child.is_file():
                    child.unlink()
                elif child.is_dir():
                    for nested in child.rglob("*"):
                        if nested.is_file():
                            nested.unlink()
                    child.rmdir()
            except OSError:
                pass
        try:
            upload_dir.rmdir()
        except OSError:
            pass

    # Remove rag slug from users.
    user_ids: List[ObjectId] = rag.get("users", [])
    for user_id in user_ids:
        doc = await find_user_by_id(str(user_id))
        if not doc:
            continue
        current_rags = [slug for slug in doc.get("rags", []) if slug != rag_slug]
        await update_user_rags(user_id, sorted(current_rags))

    # Remove rag entry and access list.
    col = get_collection()
    result = await col.delete_one({"_id": rag["_id"]})
    if result.deleted_count != 1:
        raise_http_error("DELETE_FAILED", f"Unable to delete RAG '{rag_slug}'", status.HTTP_500_INTERNAL_SERVER_ERROR)

    await write_system_log(
        event="rag.delete.complete",
        rag_slug=rag_slug,
        user_id=actor_id,
        details={"docs_deleted": len(rag.get("docs", []))},
    )
