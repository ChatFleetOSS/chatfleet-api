# generated-by: codex-agent 2026-05-04T00:00:00Z
from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document as DocxDocument
from fastapi import UploadFile
from odf.draw import Frame, Page, TextBox
from odf.opendocument import (
    OpenDocumentPresentation,
    OpenDocumentSpreadsheet,
    OpenDocumentText,
)
from odf.style import MasterPage, PageLayout, PageLayoutProperties
from odf.table import Table, TableCell, TableRow
from odf.text import P

from app.services import rags


PDF_BYTES = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>
endobj
4 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
5 0 obj
<< /Length 57 >>
stream
BT
/F1 24 Tf
100 700 Td
(ChatFleet PDF regression text) Tj
ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000241 00000 n 
0000000311 00000 n 
trailer
<< /Root 1 0 R /Size 6 >>
startxref
418
%%EOF
"""

HISTORICAL_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}

OPENDOCUMENT_EXTENSIONS = {
    ".odt": "application/vnd.oasis.opendocument.text",
    ".ods": "application/vnd.oasis.opendocument.spreadsheet",
    ".odp": "application/vnd.oasis.opendocument.presentation",
}


def _page_text(pages: list) -> str:
    return "\n".join(page.text for page in pages)


def _make_odt(path: Path) -> None:
    doc = OpenDocumentText()
    doc.text.addElement(P(text="ChatFleet ODT paragraph extraction text."))

    table = Table(name="ODT Table")
    row = TableRow()
    cell = TableCell()
    cell.addElement(P(text="ChatFleet ODT table cell"))
    row.addElement(cell)
    table.addElement(row)
    doc.text.addElement(table)
    doc.save(str(path))


def _make_ods(path: Path) -> None:
    doc = OpenDocumentSpreadsheet()
    table = Table(name="Sheet1")
    row = TableRow()
    for text in ("ChatFleet ODS spreadsheet cell", "OpenDocument spreadsheet"):
        cell = TableCell()
        cell.addElement(P(text=text))
        row.addElement(cell)
    table.addElement(row)
    doc.spreadsheet.addElement(table)
    doc.save(str(path))


def _make_odp(path: Path) -> None:
    doc = OpenDocumentPresentation()
    layout = PageLayout(name="pm1")
    layout.addElement(
        PageLayoutProperties(
            margin="0cm",
            pagewidth="28cm",
            pageheight="21cm",
            printorientation="landscape",
        )
    )
    doc.automaticstyles.addElement(layout)
    master = MasterPage(name="Default", pagelayoutname=layout)
    doc.masterstyles.addElement(master)

    slide = Page(masterpagename=master)
    frame = Frame(width="20cm", height="5cm", x="1cm", y="1cm")
    box = TextBox()
    box.addElement(P(text="ChatFleet ODP presentation slide text."))
    frame.addElement(box)
    slide.addElement(frame)
    doc.presentation.addElement(slide)
    doc.save(str(path))


async def _extract_opendocument(path: Path, extension: str) -> list:
    candidates = [
        f"_extract_{extension.removeprefix('.')}_text",
        "_extract_opendocument_text",
        "_extract_odf_text",
    ]
    if extension == ".odt":
        candidates.append("_extract_odt_text")
    for name in candidates:
        extractor = getattr(rags, name, None)
        if extractor is None:
            continue
        try:
            return await extractor(str(path))
        except TypeError:
            return await extractor(str(path), OPENDOCUMENT_EXTENSIONS[extension])
    raise AssertionError(f"No OpenDocument extractor available for {extension}")


class OpenDocumentIngestionUnitTest(unittest.IsolatedAsyncioTestCase):
    def test_allowed_extensions_cover_historical_and_opendocument_formats(self) -> None:
        expected = {**HISTORICAL_EXTENSIONS, **OPENDOCUMENT_EXTENSIONS}
        for extension, mime in expected.items():
            with self.subTest(extension=extension):
                self.assertEqual(rags.ALLOWED_EXTENSIONS.get(extension), mime)

    async def test_odt_ods_odp_extractors_read_text(self) -> None:
        makers = {".odt": _make_odt, ".ods": _make_ods, ".odp": _make_odp}
        expected_text = {
            ".odt": ["ChatFleet ODT paragraph", "ChatFleet ODT table cell"],
            ".ods": ["ChatFleet ODS spreadsheet cell", "OpenDocument spreadsheet"],
            ".odp": ["ChatFleet ODP presentation slide text"],
        }
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            for extension, maker in makers.items():
                with self.subTest(extension=extension):
                    path = root / f"sample{extension}"
                    maker(path)
                    pages = await _extract_opendocument(path, extension)
                    text = _page_text(pages)
                    for expected in expected_text[extension]:
                        self.assertIn(expected, text)

    async def test_historical_pdf_docx_txt_extractors_read_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)

            pdf_path = root / "sample.pdf"
            pdf_path.write_bytes(PDF_BYTES)
            pdf_pages = await rags._extract_pdf_text(str(pdf_path))
            self.assertIn("ChatFleet PDF regression text", _page_text(pdf_pages))

            docx_path = root / "sample.docx"
            docx = DocxDocument()
            docx.add_paragraph("ChatFleet DOCX regression paragraph.")
            table = docx.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "DOCX table left"
            table.cell(0, 1).text = "DOCX table right"
            docx.save(str(docx_path))
            docx_pages = await rags._extract_docx_text(str(docx_path))
            docx_text = _page_text(docx_pages)
            self.assertIn("ChatFleet DOCX regression paragraph", docx_text)
            self.assertIn("DOCX table left | DOCX table right", docx_text)

            txt_path = root / "sample.txt"
            txt_path.write_text("ChatFleet TXT regression text.\nSecond line.", encoding="utf-8")
            txt_pages = await rags._extract_txt_text(str(txt_path))
            self.assertIn("ChatFleet TXT regression text", _page_text(txt_pages))

    async def test_persist_file_rejects_fake_pdf_upload(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            upload_root = Path(tmp)
            upload = UploadFile(file=io.BytesIO(b"not a pdf"), filename="fake.pdf")
            with patch.object(
                rags,
                "get_runtime_overrides_sync",
                return_value=(upload_root, upload_root, 50, 0.0, 5),
            ):
                with self.assertRaisesRegex(ValueError, "PDF"):
                    await rags._persist_file("security", upload, "doc-id", ".pdf")
            self.assertFalse(list(upload_root.rglob("doc-id.pdf")))

    async def test_persist_file_rejects_fake_opendocument_uploads(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            upload_root = Path(tmp)
            for extension in OPENDOCUMENT_EXTENSIONS:
                with self.subTest(extension=extension):
                    upload = UploadFile(
                        file=io.BytesIO(b"not an opendocument package"),
                        filename=f"fake{extension}",
                    )
                    with patch.object(
                        rags,
                        "get_runtime_overrides_sync",
                        return_value=(upload_root, upload_root, 50, 0.0, 5),
                    ):
                        with self.assertRaisesRegex(ValueError, "OpenDocument|ODF|package|zip"):
                            await rags._persist_file("security", upload, f"doc-id{extension}", extension)
                    self.assertFalse(list(upload_root.rglob(f"doc-id{extension}{extension}")))

    async def test_corrupt_opendocument_packages_are_not_extractable(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            for extension in OPENDOCUMENT_EXTENSIONS:
                with self.subTest(extension=extension):
                    path = root / f"corrupt{extension}"
                    path.write_bytes(b"PK\x03\x04corrupt zip payload")
                    with self.assertRaises(Exception):
                        await _extract_opendocument(path, extension)


if __name__ == "__main__":
    unittest.main()
